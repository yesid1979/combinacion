from docx import Document
import os

filepath = r"c:\Users\yesid.piedrahita\Documents\NetBeansProjects\combinacion\plantillas\INFORME_GESTION_TEMPLATE.docx"
doc = Document(filepath)

def replace_text(doc_obj, old_text, new_text):
    for paragraph in doc_obj.paragraphs:
        if old_text in paragraph.text:
            # Fallback: replace the whole paragraph text, which might lose inline formatting
            # but is safest for matching "Fecha del Informe"
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, old_text, new_text)

replace_text(doc, 'Fecha del Informe', 'Fecha de Suscripción del Informe')

doc.save(filepath)
print("Updated successfully")
